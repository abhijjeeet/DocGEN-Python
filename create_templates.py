# create_templates.py

import os
from docx import Document

TPL_DIR = os.path.join('app', 'doc_templates')
os.makedirs(TPL_DIR, exist_ok=True)

def create_defaults():
    def write_docx(name, lines):
        path = os.path.join(TPL_DIR, f'{name}_template.docx')
        if not os.path.exists(path):
            doc = Document()
            for text in lines:
                if text.startswith('# '):
                    doc.add_heading(text[2:], level=1)
                else:
                    doc.add_paragraph(text)
            doc.save(path)

    def write_txt(name, txt):
        path = os.path.join(TPL_DIR, f'{name}_template.txt')
        if not os.path.exists(path):
            with open(path, 'w', encoding='utf-8') as f:
                f.write(txt)

    write_docx('nda', [
        '# NON-DISCLOSURE AGREEMENT',
        'Date: {{effective_date}}',
        'Party: {{party_name}}',
    ])
    write_txt('nda', """\
NON-DISCLOSURE AGREEMENT

Date: {{effective_date}}
Party: {{party_name}}
""")

    write_docx('rent', [
        '# RENT AGREEMENT',
        'Start Date: {{start_date}}',
        'Landlord: {{landlord}}',
        'Tenant: {{tenant}}',
    ])
    write_txt('rent', """\
RENT AGREEMENT

Start Date: {{start_date}}
Landlord: {{landlord}}
Tenant: {{tenant}}
""")
