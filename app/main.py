import os
import time
from flask import (
    Blueprint, render_template, url_for, request,
    redirect, flash, current_app
)
from flask_login import login_required, current_user
from docxtpl import DocxTemplate
from docx import Document

# ← Add this import so `mm` is defined




from flask_wtf import FlaskForm
from wtforms import StringField, DateField, SubmitField
from wtforms.validators import DataRequired
from fpdf import FPDF
from docx import Document
from app.models import Template
from app import db


import  tempfile, zipfile, shutil
import time
from reportlab.platypus import SimpleDocTemplate, Preformatted, Spacer
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle

main_bp = Blueprint('main', __name__)


def clean_docx_template(filename: str):
    """
    Re-save the .docx to strip hidden cruft (tracked changes, comments).
    """
    tpl_path = os.path.join(current_app.root_path, 'doc_templates', filename)
    doc = Document(tpl_path)
    doc.save(tpl_path)


def make_form_class(tmpl: Template):
    attrs = {}
    for fld in tmpl.fields:
        name = fld['name']
        label = fld['label']
        cls = DateField if fld.get('type') == 'DateField' or 'date' in name.lower() else StringField
        attrs[name] = cls(label, validators=[DataRequired()])
    attrs['submit'] = SubmitField(f"Generate {tmpl.name}")
    return type(f"{tmpl.slug.title()}Form", (FlaskForm,), attrs)


# def generate_pdf_from_docx(docx_path: str, pdf_path: str):
#     """
#     Read a .docx with python-docx and write a simple PDF via fpdf2.
#     """
#     # 1) Load the Word file
#     doc = Document(docx_path)

#     # 2) Create PDF, set margins & default font
#     pdf = FPDF(orientation='P', unit='mm', format='A4')
#     pdf.set_auto_page_break(auto=True, margin=15)
#     pdf.add_page()
#     pdf.set_font("Arial", size=12)

#     # 3) Walk through paragraphs and write them
#     for para in doc.paragraphs:
#         text = para.text.strip()
#         if not text:
#             # blank line → small vertical gap
#             pdf.ln(6)
#             continue

#         # multi_cell wraps long text automatically
#         pdf.multi_cell(w=0, h=8, txt=text)

#     # 4) Save the PDF
#     pdf.output(pdf_path)

def generate_pdf_from_docx(docx_path: str, pdf_path: str):
    """
    Reads the rendered DOCX, preserves whitespace/tabs/newlines,
    and emits a PDF via ReportLab.
    """
    # 1) Extract DOCX contents
    doc = Document(docx_path)
    story = []
    styles = getSampleStyleSheet()
    normal = styles['Normal']
    pre = ParagraphStyle(
        'Pre', parent=normal,
        fontName='Courier',    # monospaced so tabs align
        fontSize=normal.fontSize,
        leading=normal.leading,
        splitLongWords=False,
        keepLeadingSpaces=True
    )

    # 2) Turn each paragraph into a Preformatted block
    for para in doc.paragraphs:
        raw = ''.join(run.text for run in para.runs)
        raw = raw.expandtabs(8)  # or whatever tab‐width you like
        story.append(Preformatted(raw, pre))
        story.append(Spacer(1, normal.leading))

    # 3) Build the PDF
    doc_pdf = SimpleDocTemplate(
        pdf_path,
        pagesize=A4,
        leftMargin=40, rightMargin=40,
        topMargin=40, bottomMargin=40
    )
    doc_pdf.build(story)


@main_bp.route('/')
def home():
    return redirect(url_for('auth.login'))


@main_bp.route('/dashboard')
@login_required
def dashboard():
    # pull distinct category names
    rows = db.session.query(Template.category) \
                    .distinct() \
                    .order_by(Template.category) \
                    .all()
    categories = [r[0] for r in rows]

    # figure out if user clicked a category
    category = request.args.get('category')
    if category:
        templates = Template.query.filter_by(category=category).all()
    else:
        templates = []

    return render_template(
        'dashboard.html',
        categories=categories,
        category=category,
        templates=templates
    )

@main_bp.route('/generate/<slug>', methods=['GET', 'POST'])
@login_required
def generate(slug):
    tmpl = Template.query.filter_by(slug=slug).first_or_404()
    # if not current_user.can_generate():
    #     flash("You've reached your 3-document daily limit. Try again tomorrow.")
    #     return redirect(url_for('main.dashboard'))

    # 1) Clean template to strip cruft
    clean_docx_template(tmpl.docx_path)

    # 2) Build dynamic WTForm
    FormClass = make_form_class(tmpl)
    form = FormClass()

    if form.validate_on_submit():
        # 3) Gather context from form
        ctx = {field['name']: getattr(form, field['name']).data for field in tmpl.fields}

        # 4) Ensure output directory exists
        out_dir = os.path.join(current_app.root_path, '..', 'static', 'generated')
        os.makedirs(out_dir, exist_ok=True)

        # 5) Unique filename base to prevent collisions
        timestamp = int(time.time())
        filename_base = f"{current_user.id}_{slug}_{timestamp}"

        # 6) Render and save DOCX
        docx_src = os.path.join(current_app.root_path, 'doc_templates', tmpl.docx_path)
        tpl = DocxTemplate(docx_src)
        tpl.render(ctx)
        out_docx = os.path.join(out_dir, f"{filename_base}.docx")
        tpl.save(out_docx)

        # 7) Convert DOCX → PDF preserving exact formatting
        out_pdf = os.path.join(out_dir, f"{filename_base}.pdf")
        generate_pdf_from_docx(out_docx, out_pdf)






        # 8) Extract TXT from DOCX verbatim
        doc = Document(out_docx)
        out_txt = os.path.join(out_dir, f"{filename_base}.txt")
        with open(out_txt, 'w', encoding='utf-8') as f:
            for para in doc.paragraphs:
                text = ''.join(run.text for run in para.runs)
                f.write(text + '\n')

        # 9) Record generation and show downloads
        current_user.record_generation()
        return render_template('download.html',
            docx=url_for('static', filename=f"generated/{filename_base}.docx"),
            pdf=url_for('static', filename=f"generated/{filename_base}.pdf"),
            txt=url_for('static', filename=f"generated/{filename_base}.txt"),
        )

    # GET or invalid form
    return render_template('forms/template_form.html', form=form, title=tmpl.name,tmpl=tmpl)


@main_bp.route('/my-documents')
@login_required
def my_documents():
    """Show list of all docs this user has generated."""
    gen_dir = os.path.join(current_app.root_path, '..', 'static', 'generated')
    files = []
    try:
        files = os.listdir(gen_dir)
    except FileNotFoundError:
        pass

    prefix = f"{current_user.id}_"
    docs = {}
    for fn in files:
        if not fn.startswith(prefix):
            continue
        base, ext = os.path.splitext(fn)
        docs.setdefault(base, {})[ext.lstrip('.')] = fn

    # docs: { "2_nda_1612345678": {"docx": "...", "pdf": "...", "txt": "..."} }
    return render_template('my_documents.html', docs=docs)


@main_bp.route('/delete-document', methods=['POST'])
@login_required
def delete_document():
    base = request.form.get('base')
    if not base:
        flash("Invalid request.", "warning")
        return redirect(url_for('main.my_documents'))

    out_dir = os.path.join(current_app.root_path, '..', 'static', 'generated')
    removed = 0
    for ext in ('docx','pdf','txt'):
        path = os.path.join(out_dir, f"{base}.{ext}")
        if os.path.exists(path):
            try:
                os.remove(path)
                removed += 1
            except Exception as e:
                current_app.logger.error(f"Failed to delete {path}: {e}")

    if removed:
        flash(f"Deleted {removed} file(s) for document {base}.", "success")
    else:
        flash("No files were found to delete.", "info")

    return redirect(url_for('main.my_documents'))
    
@main_bp.route('/templates')
@login_required
def select_all():
    cats = [t.category for t in Template.query.distinct(Template.category)]
    templates = Template.query.all()
    return render_template('select_template.html',
        categories=cats,
        category=None,
        templates=templates
    )

@main_bp.route('/templates/<category>')
@login_required
def select_category(category):
    cats = [t.category for t in Template.query.distinct(Template.category)]
    templates = Template.query.filter_by(category=category).all()
    return render_template('select_template.html',
        categories=cats,
        category=category,
        templates=templates
    )

@main_bp.route('/pricing')
def pricing():
    return render_template('pricing.html')


# ==============================jinja working code below with pdf spacing problem

# import os
# from io import BytesIO
# from flask import (
#     Blueprint, render_template, url_for,
#     redirect, flash, current_app
# )
# from flask_login import login_required, current_user
# from docxtpl import DocxTemplate
# from docx import Document as DocxReader
# from xhtml2pdf import pisa
# from app.models import Template
# from flask_wtf import FlaskForm
# from wtforms import StringField, DateField, SubmitField
# from wtforms.validators import DataRequired

# main_bp = Blueprint('main', __name__)

# def make_form_class(tmpl: Template):
#     attrs = {}
#     for fld in tmpl.fields:
#         name, label = fld['name'], fld['label']
#         cls = DateField if fld.get('type')=='DateField' or 'date' in name.lower() else StringField
#         attrs[name] = cls(label, validators=[DataRequired()])
#     attrs['submit'] = SubmitField(f"Generate {tmpl.name}")
#     return type(f"{tmpl.slug.title()}Form", (FlaskForm,), attrs)

# @main_bp.route('/')
# def home():
#     return redirect(url_for('auth.login'))

# @main_bp.route('/dashboard')
# @login_required
# def dashboard():
#     templates = Template.query.all()
#     return render_template('dashboard.html', templates=templates)

# @main_bp.route('/generate/<slug>', methods=['GET','POST'])
# @login_required
# def generate(slug):
#     tmpl = Template.query.filter_by(slug=slug).first_or_404()
#     # if not current_user.can_generate():
#     #     flash("You've reached your 3-document daily limit. Try again tomorrow.")
#     #     return redirect(url_for('main.dashboard'))

#     # dynamically build form
#     FormClass = make_form_class(tmpl)
#     form = FormClass()

#     if form.validate_on_submit():
#         # 1) Render DOCX
#         ctx = {fld['name']: getattr(form, fld['name']).data for fld in tmpl.fields}
#         docx_src = os.path.join(current_app.root_path, 'doc_templates', tmpl.docx_path)
#         tpl = DocxTemplate(docx_src)
#         tpl.render(ctx)

#         out_dir = os.path.join(current_app.root_path, '..', 'static', 'generated')
#         os.makedirs(out_dir, exist_ok=True)
#         out_docx = os.path.join(out_dir, f"{current_user.id}_{slug}.docx")
#         tpl.save(out_docx)

#         # 2) Extract paragraphs from the generated DOCX
#         doc = DocxReader(out_docx)
#         paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

#         # 3) Save TXT: paragraphs joined by double-newline
#         out_txt = out_docx.replace('.docx', '.txt')
#         with open(out_txt, 'w', encoding='utf-8') as tf: tf.write("\n\n".join(paragraphs))

#         # 4) Build simple HTML from those same paragraphs
#         html = ['<html><head><meta charset="utf-8"></head><body>']
#         for p in paragraphs:
#             html.append(f"<p>{p}</p>")
#         html.append("</body></html>")
#         html = "\n".join(html)

#         # 5) Generate PDF via xhtml2pdf
#         out_pdf = out_docx.replace('.docx', '.pdf')
#         pdf_buf = BytesIO()
#         pisa_status = pisa.CreatePDF(html, dest=pdf_buf)
#         if pisa_status.err:
#             flash("PDF generation failed.")
#         else:
#             with open(out_pdf, 'wb') as pf:
#                 pf.write(pdf_buf.getvalue())

#         # 6) Record usage and show download links
#         current_user.record_generation()
#         return render_template('download.html',
#             docx=url_for('static', filename=f"generated/{current_user.id}_{slug}.docx"),
#             pdf =url_for('static', filename=f"generated/{current_user.id}_{slug}.pdf"),
#             txt =url_for('static', filename=f"generated/{current_user.id}_{slug}.txt"),
#         )

#     # GET or errors: show the input form
#     return render_template('forms/template_form.html', form=form, title=tmpl.name)

